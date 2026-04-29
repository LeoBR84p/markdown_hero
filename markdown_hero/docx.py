"""Render Markdown to Word (.docx) using a fixed set of styles.

Public surface is :func:`word_format`. Helpers in this module operate on
plain dictionaries produced by an internal block parser; they are not
intended for external use.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .errors import MarkdownStructureError
from .extract import remove_frontmatter

# Fixed style set used to render Markdown to .docx. ``style_overrides``
# argument of :func:`word_format` shallow-merges over this dictionary.
DEFAULT_STYLES: dict[str, Any] = {
    "body_font": "Calibri",
    "code_font": "Consolas",
    "body_size": 11,
    "code_size": 10,
    "heading_sizes": {1: 18, 2: 14, 3: 12, 4: 11, 5: 11, 6: 11},
    "code_bg": "F5F5F5",
    "table_header_bg": "E7E6E6",
    "link_color": "0563C1",
    "table_caption_centered": True,
    "table_repeat_header": True,
}


def word_format(
    md: str | Path,
    output_path: str | Path,
    *,
    template: str | Path | None = None,
    style_overrides: dict[str, Any] | None = None,
) -> Path:
    """Render Markdown to a Word ``.docx`` file with a fixed style set.

    The renderer maps headings, paragraphs, lists, blockquotes, fenced
    code, inline formatting (bold, italic, strike, code, links), tables,
    captions, and horizontal rules to native Word styles. Tables repeat
    their header row across page breaks and captions are centered above
    the table.

    Args:
        md: Markdown content as text or a path to a Markdown file.
        output_path: Destination of the generated ``.docx``. Parent
            directories are created on demand.
        template: Optional path to a ``.docx`` template whose styles are
            used as the base. When omitted, a blank document is used.
        style_overrides: Mapping that shallow-merges over
            ``DEFAULT_STYLES``. Useful to customize fonts, colors, or
            sizes without supplying a full template.

    Returns:
        The ``Path`` to the written ``.docx`` file.

    Raises:
        FrontmatterError: When the input frontmatter is invalid.
        MarkdownStructureError: When the parser encounters a block kind
            it cannot render. This indicates an internal bug, not a
            problem with the user input.
    """
    text = _read_md(md)
    body, _ = remove_frontmatter(text)

    styles = dict(DEFAULT_STYLES)
    if style_overrides:
        styles.update(style_overrides)

    doc = Document(str(template)) if template else Document()
    _ensure_custom_styles(doc, styles)

    blocks = _parse_blocks(body)
    for block in blocks:
        _render_block(doc, block, styles)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _read_md(md: str | Path) -> str:
    if isinstance(md, Path):
        return md.read_text(encoding="utf-8")
    if "\n" not in md and Path(md).exists():
        return Path(md).read_text(encoding="utf-8")
    return md


# --- Parsing em blocos -------------------------------------------------------

_FENCED_RE = re.compile(r"^(```|~~~)([^\n]*)\n(.*?)^\1\s*$", re.MULTILINE | re.DOTALL)
_HEADING_RE = re.compile(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$")
_QUOTE_RE = re.compile(r"^\s{0,3}>\s?(.*)$")
_UL_RE = re.compile(r"^(\s*)([-*+])\s+(.*)$")
_OL_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
_HR_RE = re.compile(r"^\s{0,3}(?:-{3,}|\*{3,}|_{3,})\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?(?:\s*\|\s*:?-{3,}:?)*\s*\|?\s*$")
_CAPTION_RE = re.compile(r"^Table:\s*(.+)$", re.IGNORECASE)


def _parse_blocks(body: str) -> list[dict[str, Any]]:
    """Tokenize Markdown into a flat list of typed block dictionaries."""
    blocks: list[dict[str, Any]] = []
    fenced_segments: list[tuple[int, int, dict[str, Any]]] = []
    for m in _FENCED_RE.finditer(body):
        fenced_segments.append((m.start(), m.end(), {
            "kind": "code",
            "lang": m.group(2).strip() or None,
            "code": m.group(3).rstrip("\n"),
        }))

    cursor = 0
    for start, end, code_block in fenced_segments:
        if start > cursor:
            blocks.extend(_parse_non_code(body[cursor:start]))
        blocks.append(code_block)
        cursor = end
    if cursor < len(body):
        blocks.extend(_parse_non_code(body[cursor:]))
    return blocks


def _parse_non_code(text: str) -> list[dict[str, Any]]:
    """Walk the input line-by-line and emit non-code blocks.

    The function is a thin state machine that delegates each block kind
    to a focused handler. ``pending_caption`` tracks the Pandoc-style
    ``Table:`` line so it can be attached to the most recent table.
    """
    lines = text.split("\n")
    blocks: list[dict[str, Any]] = []
    state = _ParseState(lines=lines, i=0, blocks=blocks, pending_caption=None)
    while state.i < len(lines):
        line = lines[state.i]
        if not line.strip():
            state.i += 1
            continue
        if _HR_RE.match(line):
            blocks.append({"kind": "hr"})
            state.i += 1
            continue
        if _consume_heading(line, state):
            continue
        if _consume_table(line, state):
            continue
        if _consume_caption(line, state):
            continue
        if _consume_quote(line, state):
            continue
        if _consume_list(line, state):
            continue
        _consume_paragraph(line, state)
    return blocks


@dataclass
class _ParseState:
    """Mutable cursor used while walking lines in :func:`_parse_non_code`."""

    lines: list[str]
    i: int
    blocks: list[dict[str, Any]]
    pending_caption: str | None


def _consume_heading(line: str, state: _ParseState) -> bool:
    m = _HEADING_RE.match(line)
    if not m:
        return False
    state.blocks.append({"kind": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
    state.i += 1
    return True


def _consume_table(line: str, state: _ParseState) -> bool:
    if "|" not in line:
        return False
    if state.i + 1 >= len(state.lines):
        return False
    if not _TABLE_SEP_RE.match(state.lines[state.i + 1]):
        return False
    headers = _split_row(line)
    aligns = _parse_align(state.lines[state.i + 1])
    rows: list[list[str]] = []
    j = state.i + 2
    while j < len(state.lines) and state.lines[j].strip() and "|" in state.lines[j]:
        rows.append(_split_row(state.lines[j]))
        j += 1
    state.blocks.append({
        "kind": "table",
        "headers": headers,
        "rows": rows,
        "alignments": aligns,
        "caption": state.pending_caption,
    })
    state.pending_caption = None
    state.i = j
    return True


def _consume_caption(line: str, state: _ParseState) -> bool:
    cap = _CAPTION_RE.match(line.strip())
    if not cap:
        return False
    text = cap.group(1).strip()
    if state.blocks and state.blocks[-1]["kind"] == "table":
        state.blocks[-1]["caption"] = text
    else:
        state.pending_caption = text
    state.i += 1
    return True


def _consume_quote(line: str, state: _ParseState) -> bool:
    if not _QUOTE_RE.match(line):
        return False
    buf: list[str] = []
    while state.i < len(state.lines) and (
        _QUOTE_RE.match(state.lines[state.i])
        or (state.lines[state.i].strip() and not _is_block_start(state.lines[state.i]))
    ):
        qm = _QUOTE_RE.match(state.lines[state.i])
        buf.append(qm.group(1) if qm else state.lines[state.i].strip())
        state.i += 1
    state.blocks.append({"kind": "quote", "text": " ".join(buf).strip()})
    return True


def _consume_list(line: str, state: _ParseState) -> bool:
    if not (_UL_RE.match(line) or _OL_RE.match(line)):
        return False
    items: list[dict[str, Any]] = []
    while state.i < len(state.lines) and state.lines[state.i].strip() and (
        _UL_RE.match(state.lines[state.i]) or _OL_RE.match(state.lines[state.i])
    ):
        items.append(_parse_list_item(state.lines[state.i]))
        state.i += 1
    state.blocks.append({"kind": "list", "items": items})
    return True


def _parse_list_item(line: str) -> dict[str, Any]:
    ul = _UL_RE.match(line)
    if ul:
        indent, _, text_part = ul.groups()
        return {"ordered": False, "level": len(indent) // 2, "text": text_part}
    ol = _OL_RE.match(line)
    assert ol is not None  # caller already matched
    indent, _, text_part = ol.groups()
    return {"ordered": True, "level": len(indent) // 2, "text": text_part}


def _consume_paragraph(line: str, state: _ParseState) -> None:
    buf = [line]
    state.i += 1
    while (
        state.i < len(state.lines)
        and state.lines[state.i].strip()
        and not _is_block_start(state.lines[state.i])
    ):
        buf.append(state.lines[state.i])
        state.i += 1
    state.blocks.append({"kind": "paragraph", "text": " ".join(b.strip() for b in buf)})


def _is_block_start(line: str) -> bool:
    return bool(
        _HEADING_RE.match(line)
        or _UL_RE.match(line)
        or _OL_RE.match(line)
        or _QUOTE_RE.match(line)
        or _HR_RE.match(line)
    )


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _parse_align(sep: str) -> list[str]:
    out: list[str] = []
    for part in _split_row(sep):
        left = part.startswith(":")
        right = part.endswith(":")
        if left and right:
            out.append("center")
        elif right:
            out.append("right")
        elif left:
            out.append("left")
        else:
            out.append("default")
    return out


# --- Renderização ------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(?P<code>`[^`\n]+`)"
    r"|(?P<bold>\*\*[^*\n]+\*\*|__[^_\n]+__)"
    r"|(?P<italic>\*[^*\n]+\*|_[^_\n]+_)"
    r"|(?P<strike>~~[^~\n]+~~)"
    r"|(?P<image>!\[[^\]]*\]\([^)]+\))"
    r"|(?P<link>\[[^\]]+\]\([^)]+\))"
)


def _render_heading(doc: Any, block: dict[str, Any], styles: dict[str, Any]) -> None:
    p = doc.add_heading(level=min(block["level"], 9))
    p.text = ""
    _render_inline(p, block["text"], styles)


def _render_paragraph(doc: Any, block: dict[str, Any], styles: dict[str, Any]) -> None:
    p = doc.add_paragraph(style="Normal")
    _render_inline(p, block["text"], styles)


def _render_quote(doc: Any, block: dict[str, Any], styles: dict[str, Any]) -> None:
    try:
        p = doc.add_paragraph(style="Quote")
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
    _render_inline(p, block["text"], styles)
    for run in p.runs:
        run.italic = True


def _render_list(doc: Any, block: dict[str, Any], styles: dict[str, Any]) -> None:
    for item in block["items"]:
        style_name = "List Number" if item["ordered"] else "List Bullet"
        try:
            p = doc.add_paragraph(style=style_name)
        except KeyError:
            p = doc.add_paragraph()
        _render_inline(p, item["text"], styles)


def _render_code_block(doc: Any, block: dict[str, Any], styles: dict[str, Any]) -> None:
    _render_code(doc, block["code"], styles)


def _render_hr(doc: Any, _block: dict[str, Any], _styles: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


_BLOCK_RENDERERS: dict[str, Callable[[Any, dict[str, Any], dict[str, Any]], None]] = {
    "heading": _render_heading,
    "paragraph": _render_paragraph,
    "quote": _render_quote,
    "list": _render_list,
    "code": _render_code_block,
    "hr": _render_hr,
    "table": lambda doc, block, styles: _render_table(doc, block, styles),
}


def _render_block(doc: Any, block: dict[str, Any], styles: dict[str, Any]) -> None:
    renderer = _BLOCK_RENDERERS.get(block["kind"])
    if renderer is None:
        raise MarkdownStructureError(f"unknown block kind: {block['kind']!r}")
    renderer(doc, block, styles)


def _render_inline(paragraph: Any, text: str, styles: dict[str, Any]) -> None:
    """Apply inline formatting (bold, italic, code, link) to a Word paragraph."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("bold"):
            run = paragraph.add_run(m.group("bold").strip("*_"))
            run.bold = True
        elif m.group("italic"):
            run = paragraph.add_run(m.group("italic").strip("*_"))
            run.italic = True
        elif m.group("strike"):
            run = paragraph.add_run(m.group("strike").strip("~"))
            run.font.strike = True
        elif m.group("code"):
            run = paragraph.add_run(m.group("code").strip("`"))
            run.font.name = styles["code_font"]
            run.font.size = Pt(styles["code_size"])
            _shade_run(run, styles["code_bg"])
        elif m.group("image"):
            inner = m.group("image")
            mm = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", inner)
            if mm:
                run = paragraph.add_run(f"[image: {mm.group(1) or mm.group(2)}]")
                run.italic = True
        elif m.group("link"):
            inner = m.group("link")
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", inner)
            if mm:
                _add_hyperlink(paragraph, mm.group(2), mm.group(1), styles["link_color"])
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _render_code(doc: Any, code: str, styles: dict[str, Any]) -> None:
    try:
        p = doc.add_paragraph(style="Code Block")
    except KeyError:
        p = doc.add_paragraph()
    for i, line in enumerate(code.splitlines() or [""]):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = styles["code_font"]
        run.font.size = Pt(styles["code_size"])
    _shade_paragraph(p, styles["code_bg"])


def _render_table(doc: Any, block: dict[str, Any], styles: dict[str, Any]) -> None:
    headers = block["headers"]
    if not headers:
        return
    if block.get("caption") and styles["table_caption_centered"]:
        _render_table_caption(doc, block["caption"])
    rows = block["rows"]
    aligns = block.get("alignments") or ["default"] * len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    _render_table_header(table, headers, styles)
    _render_table_body(table, headers, rows, aligns, styles)


def _render_table_caption(doc: Any, caption: str) -> None:
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)


def _render_table_header(table: Any, headers: list[str], styles: dict[str, Any]) -> None:
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        cell = hdr_cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        _render_inline(p, h, styles)
        for run in p.runs:
            run.bold = True
        _set_cell_align(p, "center")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, styles["table_header_bg"])
    if styles["table_repeat_header"]:
        _repeat_table_header(table.rows[0])


def _render_table_body(
    table: Any,
    headers: list[str],
    rows: list[list[str]],
    aligns: list[str],
    styles: dict[str, Any],
) -> None:
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for j in range(len(headers)):
            text_cell = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _render_inline(p, text_cell, styles)
            align = aligns[j] if j < len(aligns) else "default"
            _set_cell_align(p, align)


def _set_cell_align(paragraph: Any, align: str) -> None:
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _shade_cell(cell: Any, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _shade_paragraph(paragraph: Any, hex_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def _shade_run(run: Any, hex_color: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    r_pr.append(shd)


def _repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_hyperlink(paragraph: Any, url: str, text: str, color: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# --- Estilos customizados ----------------------------------------------------

def _ensure_custom_styles(doc: Any, styles: dict[str, Any]) -> None:
    from docx.enum.style import WD_STYLE_TYPE

    style_collection = doc.styles
    names = {s.name for s in style_collection}

    if "Code Block" not in names:
        s = style_collection.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = styles["code_font"]
        s.font.size = Pt(styles["code_size"])

    if "Code Char" not in names:
        s = style_collection.add_style("Code Char", WD_STYLE_TYPE.CHARACTER)
        s.font.name = styles["code_font"]
        s.font.size = Pt(styles["code_size"])

    # Ajustes de fonte default (apenas se 'Normal' não foi imposta por template).
    normal = style_collection["Normal"]
    if normal.font.name in (None, "Calibri"):
        normal.font.name = styles["body_font"]
        normal.font.size = Pt(styles["body_size"])

    # Tamanhos de heading.
    for lvl, size in styles["heading_sizes"].items():
        try:
            h = style_collection[f"Heading {lvl}"]
            h.font.size = Pt(size)
            h.font.bold = True
            if h.font.color and h.font.color.rgb is None:
                h.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass
